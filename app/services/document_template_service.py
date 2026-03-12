"""
文档模板服务
"""
import os
import uuid
from datetime import datetime
from common import log_
from common.minio_client import MinioClient
from app.dao.document_template_dao import DocumentTemplateDao, TemplateGenerationDao
from app.entity.document_template import DocumentTemplate, TemplateGeneration
from app.models.llm.model_factory import get_llm_model
from app.utils.document_loader import DocumentLoader


class DocumentTemplateService:
    """文档模板服务类"""
    
    def __init__(self):
        self.template_dao = DocumentTemplateDao()
        self.generation_dao = TemplateGenerationDao()
        self.minio_client = MinioClient()
        self.document_loader = DocumentLoader()
    
    def upload_template(self, user_id, file_path, file_name, tags, file_size):
        """
        上传文档模板
        
        Args:
            user_id: 用户ID
            file_path: 本地文件路径
            file_name: 文件名
            tags: 标签
            file_size: 文件大小
        
        Returns:
            dict: 上传结果
        """
        try:
            log_.info(f"开始上传模板: user_id={user_id}, file_name={file_name}")
            
            # 确定文件类型 - 使用file_path确保获取正确的扩展名
            file_ext = os.path.splitext(file_path)[1].lower()
            log_.info(f"文件扩展名: {file_ext}")
            
            if file_ext in ['.doc', '.docx']:
                file_type = 'word'
            elif file_ext == '.pdf':
                file_type = 'pdf'
            else:
                return {'success': False, 'message': '不支持的文件格式，仅支持Word和PDF'}
            
            # 生成MinIO存储路径
            object_name = f"templates/{user_id}/{uuid.uuid4()}{file_ext}"
            log_.info(f"MinIO路径: {object_name}")
            
            # 上传到MinIO
            log_.info(f"开始上传到MinIO, 文件路径: {file_path}")
            self.minio_client.upload_file(file_path, object_name)
            log_.info("MinIO上传成功")
            
            # 保存到数据库
            log_.info("开始保存到数据库")
            template = DocumentTemplate(
                name=file_name,
                tags=tags,
                file_type=file_type,
                file_size=file_size,
                minio_path=object_name,
                created_by=user_id
            )
            template_id = self.template_dao.insert(template)
            log_.info(f"数据库保存成功, template_id={template_id}")
            
            return {
                'success': True,
                'template_id': template_id,
                'message': '上传成功'
            }
        
        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            log_.error(f"上传模板失败: {str(e)}\n详细错误:\n{error_detail}")
            return {'success': False, 'message': f'上传失败: {str(e)}'}
    
    def get_template_list(self, user_id, tag=None, search=None):
        """
        获取模板列表
        
        Args:
            user_id: 用户ID
            tag: 标签筛选
            search: 搜索关键词
        
        Returns:
            list: 模板列表
        """
        templates = self.template_dao.find_by_user(user_id, tag, search)
        return templates
    
    def get_template_detail(self, template_id, user_id):
        """
        获取模板详情
        
        Args:
            template_id: 模板ID
            user_id: 用户ID
        
        Returns:
            dict: 模板详情
        """
        template = self.template_dao.find_by_id_and_user(template_id, user_id)
        return template
    
    def delete_template(self, template_id, user_id):
        """
        删除模板
        
        Args:
            template_id: 模板ID
            user_id: 用户ID
        
        Returns:
            dict: 删除结果
        """
        try:
            # 获取模板信息
            template = self.template_dao.find_by_id_and_user(template_id, user_id)
            if not template:
                return {'success': False, 'message': '模板不存在或无权限'}
            
            # 删除MinIO文件
            try:
                self.minio_client.delete_file(template['minio_path'])
            except Exception as e:
                log_.warning(f"删除MinIO文件失败: {str(e)}")
            
            # 删除数据库记录
            success = self.template_dao.delete_by_id_and_user(template_id, user_id)
            
            if success:
                return {'success': True, 'message': '删除成功'}
            else:
                return {'success': False, 'message': '删除失败'}
        
        except Exception as e:
            log_.error(f"删除模板失败: {str(e)}")
            return {'success': False, 'message': f'删除失败: {str(e)}'}
    
    def generate_document(self, user_id, template_id, user_input):
        """
        生成文档
        
        Args:
            user_id: 用户ID
            template_id: 模板ID
            user_input: 用户输入的个人信息
        
        Returns:
            dict: 生成结果
        """
        try:
            # 获取模板信息
            template = self.template_dao.find_by_id_and_user(template_id, user_id)
            if not template:
                return {'success': False, 'message': '模板不存在或无权限'}
            
            # 创建生成记录
            generation = TemplateGeneration(
                template_id=template_id,
                user_id=user_id,
                user_input=user_input,
                status='generating'
            )
            generation_id = self.generation_dao.insert(generation)
            
            # 下载模板文件（使用系统临时目录，兼容Windows）
            import tempfile
            temp_dir = tempfile.gettempdir()
            file_ext = os.path.splitext(template['minio_path'])[1]
            safe_name = template['name'].replace(os.sep, '_')
            local_template_path = os.path.join(temp_dir, f"{uuid.uuid4()}_{safe_name}{file_ext}")
            self.minio_client.download_file(template['minio_path'], local_template_path)
            
            # 读取模板内容
            template_content = self.document_loader.load(local_template_path)
            
            # 清理临时文件
            try:
                os.remove(local_template_path)
            except:
                pass
            
            # 调用LLM生成文档
            llm = get_llm_model()
            
            prompt = f"""你是一个专业的文档生成助手。请根据以下模板格式和用户信息，生成一份完整的文档。

模板格式参考：
{template_content}

用户提供的信息：
{user_input}

要求：
1. 严格按照模板的格式和结构生成文档
2. 将用户提供的信息填充到模板相应位置
3. 保持专业、正式的文档风格
4. 如果用户信息不完整，请合理补充或标注待填写
5. 直接输出文档内容，不要包含解释说明

请生成完整的文档内容："""
            
            generated_content = llm.invoke(prompt)
            
            # 更新生成记录
            self.generation_dao.update_status(
                generation_id,
                'completed',
                generated_content=generated_content
            )
            
            return {
                'success': True,
                'generation_id': generation_id,
                'content': generated_content,
                'message': '生成成功'
            }
        
        except Exception as e:
            log_.error(f"生成文档失败: {str(e)}")
            # 更新失败状态
            if 'generation_id' in locals():
                self.generation_dao.update_status(generation_id, 'failed')
            return {'success': False, 'message': f'生成失败: {str(e)}'}
    
    def get_generation_history(self, user_id, page=1, page_size=20):
        """
        获取生成历史
        
        Args:
            user_id: 用户ID
            page: 页码
            page_size: 每页数量
        
        Returns:
            dict: 历史记录和总数
        """
        history = self.generation_dao.find_by_user(user_id, page, page_size)
        total = self.generation_dao.count_by_user(user_id)
        
        return {
            'list': history,
            'total': total,
            'page': page,
            'page_size': page_size
        }
    
    def get_generation_detail(self, generation_id, user_id):
        """
        获取生成记录详情
        
        Args:
            generation_id: 生成记录ID
            user_id: 用户ID
        
        Returns:
            dict: 生成记录详情
        """
        generation = self.generation_dao.find_by_id_and_user(generation_id, user_id)
        return generation
    
    def get_template_file_url(self, template_id, user_id):
        """
        获取模板文件下载URL
        
        Args:
            template_id: 模板ID
            user_id: 用户ID
        
        Returns:
            str: 下载URL
        """
        template = self.template_dao.find_by_id_and_user(template_id, user_id)
        if not template:
            return None
        
        return self.minio_client.get_presigned_url(template['minio_path'])
    
    def export_generation(self, generation_id, user_id, export_format='word'):
        """
        导出生成的文档为Word或PDF格式
        
        Args:
            generation_id: 生成记录ID
            user_id: 用户ID
            export_format: 导出格式 ('word' 或 'pdf')
        
        Returns:
            dict: 包含文件路径和文件名的字典
        """
        try:
            # 获取生成记录
            generation = self.generation_dao.find_by_id_and_user(generation_id, user_id)
            if not generation:
                return {'success': False, 'message': '记录不存在或无权限'}
            
            if generation['status'] != 'completed':
                return {'success': False, 'message': '文档尚未生成完成'}
            
            content = generation['generated_content']
            template_name = generation.get('template_name', '文档')
            
            # 生成文件（使用系统临时目录，兼容Windows）
            import tempfile
            temp_dir = tempfile.gettempdir()
            os.makedirs(temp_dir, exist_ok=True)
            
            if export_format == 'word':
                # 生成Word文档
                file_path = self._generate_word(content, template_name, temp_dir)
                file_name = f"{template_name}.docx"
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif export_format == 'pdf':
                # 生成PDF文档
                file_path = self._generate_pdf(content, template_name, temp_dir)
                file_name = f"{template_name}.pdf"
                content_type = 'application/pdf'
            else:
                return {'success': False, 'message': '不支持的导出格式'}
            
            return {
                'success': True,
                'file_path': file_path,
                'file_name': file_name,
                'content_type': content_type
            }
        
        except Exception as e:
            log_.error(f"导出文档失败: {str(e)}")
            return {'success': False, 'message': f'导出失败: {str(e)}'}
    
    def _generate_word(self, content, file_name, temp_dir):
        """
        生成Word文档
        
        Args:
            content: 文档内容
            file_name: 文件名
            temp_dir: 临时目录
        
        Returns:
            str: 文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx库未安装，请使用 'pip install python-docx' 安装")
        
        doc = Document()
        
        # 设置标题
        title = doc.add_heading(file_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加内容，处理换行
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text)
                para.paragraph_format.first_line_indent = Inches(0.5)
        
        # 保存文件
        file_path = os.path.join(temp_dir, f"{uuid.uuid4()}_{file_name}.docx")
        doc.save(file_path)
        
        return file_path
    
    def _generate_pdf(self, content, file_name, temp_dir):
        """
        生成PDF文档
        
        Args:
            content: 文档内容
            file_name: 文件名
            temp_dir: 临时目录
        
        Returns:
            str: 文件路径
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.units import inch
            
            # 尝试注册中文字体
            try:
                # Windows系统字体路径
                font_paths = [
                    'C:/Windows/Fonts/simsun.ttc',
                    'C:/Windows/Fonts/msyh.ttc',
                    'C:/Windows/Fonts/simhei.ttf',
                    '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                    '/usr/share/fonts/truetype/arphic/uming.ttc'
                ]
                font_registered = False
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        font_registered = True
                        break
                
                if not font_registered:
                    # 如果没有找到中文字体，使用默认字体
                    chinese_font = 'Helvetica'
                else:
                    chinese_font = 'ChineseFont'
            except:
                chinese_font = 'Helvetica'
            
            file_path = os.path.join(temp_dir, f"{uuid.uuid4()}_{file_name}.pdf")
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            
            styles = getSampleStyleSheet()
            # 创建中文样式
            chinese_style = ParagraphStyle(
                'ChineseStyle',
                parent=styles['Normal'],
                fontName=chinese_font,
                fontSize=12,
                leading=20,
                firstLineIndent=24
            )
            
            title_style = ParagraphStyle(
                'TitleStyle',
                parent=styles['Title'],
                fontName=chinese_font,
                fontSize=18,
                alignment=1  # 居中
            )
            
            story = []
            
            # 添加标题
            story.append(Paragraph(file_name, title_style))
            story.append(Spacer(1, 0.5*inch))
            
            # 添加内容
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # 处理特殊字符
                    safe_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe_text, chinese_style))
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            return file_path
            
        except ImportError:
            raise ImportError("reportlab库未安装，请使用 'pip install reportlab' 安装")
