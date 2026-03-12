"""
文档生成服务 - 处理AI文档生成和格式转换
"""
import os
import uuid
import tempfile
from werkzeug.utils import secure_filename
from datetime import datetime
from common import log_
from common.minio_client import MinioClient
from app.models.llm.ollama_model import OllamaModel
from app.entity.doc_generation import DocGeneration
from app.dao.doc_generation_dao import DocGenerationDAO
from app.dao.template_dao import TemplateDAO
from common.error_codes import APIException, ErrorCode
from docx import Document
from docx.shared import Inches
import fitz

class DocGenerationService:
    """文档生成服务类"""
    
    def __init__(self):
        """初始化文档生成服务"""
        self.minio_client = MinioClient.get_instance()
        self.llm_model = OllamaModel()
        self.doc_generation_dao = DocGenerationDAO()
        self.template_dao = TemplateDAO()
    
    def generate_document(self, user_id, template_id, user_input):
        """生成文档"""
        try:
            template = self.template_dao.select_one({'id': template_id})
            if not template:
                raise APIException(ErrorCode.VALIDATION_ERROR, msg="模板不存在")
            
            template_content = ""
            if hasattr(template, 'minio_path') and template.minio_path:
                try:
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.tmp')
                    temp_file.close()
                    self.minio_client.download_file(template.minio_path, temp_file.name)
                    template_content = self._extract_file_content(temp_file.name)
                    os.unlink(temp_file.name)
                except Exception as e:
                    log_.warning(f"提取模板内容失败: {str(e)}")
            
            prompt = self._build_generation_prompt(template_content, user_input)
            generated_content = self.llm_model.invoke(prompt)
            
            doc_generation = DocGeneration(
                user_id=user_id,
                template_id=template_id,
                user_input=user_input,
                generated_content=generated_content,
                status='completed'
            )
            
            result = self.doc_generation_dao.create(doc_generation)
            result_id = result.get('id') if isinstance(result, dict) else getattr(result, 'id', None)
            
            word_path = self._generate_word_file(result_id, generated_content, user_id)
            pdf_path = self._generate_pdf_file(result_id, generated_content, user_id)
            
            if word_path or pdf_path:
                self.doc_generation_dao.update_result(result_id, word_minio_path=word_path, pdf_minio_path=pdf_path)
            
            return {
                'id': result_id,
                'generated_content': generated_content,
                'word_minio_path': word_path,
                'pdf_minio_path': pdf_path
            }
            
        except APIException:
            raise
        except Exception as e:
            log_.error(f"生成文档失败: {str(e)}")
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"生成文档失败: {str(e)}")
    
    def _extract_file_content(self, file_path):
        """提取文件内容"""
        try:
            if file_path.lower().endswith('.pdf'):
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            elif file_path.lower().endswith(('.doc', '.docx')):
                doc = Document(file_path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            else:
                return ""
        except Exception as e:
            log_.error(f"提取文件内容失败: {str(e)}")
            return ""
    
    def _build_generation_prompt(self, template_content, user_input):
        """构建生成提示词"""
        prompt = f"""你是一个专业的文档生成AI助手。请根据以下要求生成文档：

用户输入信息：
{user_input}

"""
        if template_content and template_content.strip():
            prompt += f"""参考模板内容：
{template_content}

请基于上述模板结构和用户提供的信息，生成一份完整的文档。
"""
        else:
            prompt += """请根据用户提供的信息，生成一份格式规范、内容完整的文档。
"""
        
        prompt += """
要求：
1. 保持语言正式、专业
2. 内容完整，逻辑清晰
3. 格式规范，段落分明
4. 直接返回文档内容，不要添加任何额外说明
"""
        return prompt
    
    def _generate_word_file(self, result_id, content, user_id):
        """生成Word文件"""
        try:
            temp_dir = os.path.abspath("./data/temp")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file_path = os.path.join(temp_dir, f"{uuid.uuid4().hex}.docx")
            
            doc = Document()
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
                else:
                    doc.add_paragraph()
            
            doc.save(temp_file_path)
            
            if os.path.exists(temp_file_path):
                minio_path = f"generated_docs/{user_id}/{result_id}.docx"
                self.minio_client.upload_file(temp_file_path, minio_path)
                os.remove(temp_file_path)
                return minio_path
            return None
        except Exception as e:
            log_.error(f"生成Word文件失败: {str(e)}")
            return None
    
    def _generate_pdf_file(self, result_id, content, user_id):
        """生成PDF文件"""
        try:
            temp_dir = os.path.abspath("./data/temp")
            os.makedirs(temp_dir, exist_ok=True)
            word_temp = os.path.join(temp_dir, f"{uuid.uuid4().hex}.docx")
            pdf_temp = os.path.join(temp_dir, f"{uuid.uuid4().hex}.pdf")
            
            doc = Document()
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
                else:
                    doc.add_paragraph()
            doc.save(word_temp)
            
            try:
                import subprocess
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, word_temp],
                    capture_output=True,
                    timeout=30
                )
                if result.returncode == 0 and os.path.exists(pdf_temp):
                    minio_path = f"generated_docs/{user_id}/{result_id}.pdf"
                    self.minio_client.upload_file(pdf_temp, minio_path)
                    if os.path.exists(word_temp):
                        os.remove(word_temp)
                    if os.path.exists(pdf_temp):
                        os.remove(pdf_temp)
                    return minio_path
            except Exception as e:
                log_.warning(f"LibreOffice转换失败: {str(e)}")
            
            if os.path.exists(word_temp):
                os.remove(word_temp)
            if os.path.exists(pdf_temp):
                os.remove(pdf_temp)
            
            return None
        except Exception as e:
            log_.error(f"生成PDF文件失败: {str(e)}")
            return None
    
    def get_history(self, user_id, limit=100, offset=0):
        """获取生成历史"""
        try:
            history = self.doc_generation_dao.find_all_by_user(user_id, limit, offset)
            for item in history:
                if item.get('word_minio_path'):
                    item['word_url'] = self.minio_client.get_presigned_url(item['word_minio_path'])
                if item.get('pdf_minio_path'):
                    item['pdf_url'] = self.minio_client.get_presigned_url(item['pdf_minio_path'])
            return history
        except Exception as e:
            log_.error(f"获取历史记录失败: {str(e)}")
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"获取历史记录失败: {str(e)}")
    
    def get_result_detail(self, result_id, user_id):
        """获取生成结果详情"""
        try:
            result = self.doc_generation_dao.select_one({'id': result_id, 'user_id': user_id})
            if not result:
                raise FileNotFoundError("结果不存在")
            
            result_dict = result.to_dict() if hasattr(result, 'to_dict') else dict(result)
            
            if result_dict.get('word_minio_path'):
                result_dict['word_url'] = self.minio_client.get_presigned_url(result_dict['word_minio_path'])
            if result_dict.get('pdf_minio_path'):
                result_dict['pdf_url'] = self.minio_client.get_presigned_url(result_dict['pdf_minio_path'])
            
            return result_dict
        except FileNotFoundError:
            raise
        except Exception as e:
            log_.error(f"获取结果详情失败: {str(e)}")
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"获取结果详情失败: {str(e)}")
    
    def download_document(self, result_id, user_id, format='word'):
        """下载生成的文档"""
        try:
            result = self.doc_generation_dao.select_one({'id': result_id, 'user_id': user_id})
            if not result:
                raise FileNotFoundError("文档不存在")
            
            result_dict = result.to_dict() if hasattr(result, 'to_dict') else dict(result)
            
            if format == 'word' and result_dict.get('word_minio_path'):
                return self.minio_client.get_presigned_url(result_dict['word_minio_path'])
            elif format == 'pdf' and result_dict.get('pdf_minio_path'):
                return self.minio_client.get_presigned_url(result_dict['pdf_minio_path'])
            else:
                raise APIException(ErrorCode.VALIDATION_ERROR, msg=f"不支持的格式或文件不存在: {format}")
        except FileNotFoundError:
            raise
        except Exception as e:
            log_.error(f"下载文档失败: {str(e)}")
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"下载文档失败: {str(e)}")
