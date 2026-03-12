"""
模板文档服务层
"""
import io
import os
import tempfile
from datetime import datetime

from app.dao.template_dao import DocumentTemplateDAO, GeneratedDocumentDAO
from app.entity.template import DocumentTemplate, GeneratedDocument
from app.models.llm.ollama_model import OllamaModel
from common.minio_client import MinioClient
from common import log_
from common.error_codes import APIException, ErrorCode


class TemplateService:
    """模板文档服务"""
    
    TEMPLATE_CATEGORIES = ['简历', '论文', '报告', '合同', '信函', '其他']
    
    def __init__(self):
        try:
            self.ollama_model = OllamaModel()
        except Exception as e:
            log_.warning(f"初始化Ollama模型失败: {str(e)}")
            self.ollama_model = None
        
        try:
            self.minio_client = MinioClient.get_instance()
        except Exception as e:
            log_.warning(f"初始化MinIO客户端失败: {str(e)}")
            self.minio_client = None
    
    def upload_template(self, file_object, name, category, description, user_id, is_public=False):
        """上传模板文件"""
        try:
            user_id = int(user_id)
            filename = getattr(file_object, 'filename', getattr(file_object, 'name', 'template'))
            file_type = self._get_file_type(filename)
            file_size = self._get_file_size(file_object)
            
            minio_path = f"templates/{user_id}/{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
            
            file_data = file_object.read()
            
            if self.minio_client:
                self.minio_client.upload_bytes(
                    data=file_data,
                    object_name=minio_path,
                    content_type=self._get_content_type(file_type)
                )
            else:
                log_.warning("MinIO客户端未初始化，文件将不会上传到存储")
            
            template = DocumentTemplate(
                name=name,
                description=description,
                category=category,
                minio_path=minio_path,
                file_type=file_type,
                file_size=file_size,
                created_by=user_id,
                is_public=is_public
            )
            
            result = DocumentTemplateDAO.create(template)
            return result
            
        except APIException:
            raise
        except Exception as e:
            log_.error(f"上传模板失败: {str(e)}", exc_info=True)
            raise APIException(ErrorCode.FILE_UPLOAD_ERROR, msg=f"上传模板失败: {str(e)}")
    
    def get_template_list(self, user_id, category=None, keyword=None):
        """获取模板列表"""
        try:
            user_id = int(user_id)
            if keyword:
                return DocumentTemplateDAO.search_templates(keyword, user_id)
            elif category:
                return DocumentTemplateDAO.find_by_category(category, user_id)
            else:
                return DocumentTemplateDAO.find_by_user_id(user_id)
        except APIException:
            raise
        except Exception as e:
            log_.error(f"获取模板列表失败: {str(e)}", exc_info=True)
            raise APIException(ErrorCode.DATABASE_ERROR, msg=f"获取模板列表失败: {str(e)}")
    
    def get_template_file(self, template_id, user_id):
        """获取模板文件内容"""
        try:
            user_id = int(user_id)
            template_id = int(template_id)
            
            template = DocumentTemplateDAO.find_by_id(template_id)
            if not template:
                raise APIException(ErrorCode.RESOURCE_NOT_FOUND, msg="模板不存在")
            
            if template['created_by'] != user_id and not template['is_public']:
                raise APIException(ErrorCode.UNAUTHORIZED, msg="无权访问此模板")
            
            if not self.minio_client:
                raise APIException(ErrorCode.SYSTEM_ERROR, msg="存储服务未初始化")
            
            file_data = self.minio_client.download_file(template['minio_path'])
            return file_data, template['file_type'], template['name']
            
        except APIException:
            raise
        except Exception as e:
            log_.error(f"获取模板文件失败: {str(e)}", exc_info=True)
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"获取模板文件失败: {str(e)}")
    
    def delete_template(self, template_id, user_id):
        """删除模板"""
        try:
            user_id = int(user_id)
            template_id = int(template_id)
            
            template = DocumentTemplateDAO.find_by_id(template_id)
            if not template:
                raise APIException(ErrorCode.RESOURCE_NOT_FOUND, msg="模板不存在")
            
            if template['created_by'] != user_id:
                raise APIException(ErrorCode.UNAUTHORIZED, msg="无权删除此模板")
            
            if template['minio_path'] and self.minio_client:
                try:
                    self.minio_client.delete_file(template['minio_path'])
                except Exception as e:
                    log_.warning(f"删除MinIO文件失败: {str(e)}")
            
            return DocumentTemplateDAO.delete(template_id)
            
        except APIException:
            raise
        except Exception as e:
            log_.error(f"删除模板失败: {str(e)}", exc_info=True)
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"删除模板失败: {str(e)}")
    
    def get_categories(self, user_id=None):
        """获取所有分类"""
        try:
            if user_id:
                user_id = int(user_id)
            categories = DocumentTemplateDAO.get_all_categories(user_id)
            all_categories = list(set(self.TEMPLATE_CATEGORIES + categories))
            return sorted(all_categories)
        except Exception as e:
            log_.error(f"获取分类失败: {str(e)}")
            return self.TEMPLATE_CATEGORIES
    
    def generate_document(self, template_id, user_input, output_format, user_id):
        """生成文档"""
        try:
            user_id = int(user_id)
            template_id = int(template_id)
            
            template = DocumentTemplateDAO.find_by_id(template_id)
            if not template:
                raise APIException(ErrorCode.RESOURCE_NOT_FOUND, msg="模板不存在")
            
            if template['created_by'] != user_id and not template['is_public']:
                raise APIException(ErrorCode.UNAUTHORIZED, msg="无权使用此模板")
            
            template_text = ""
            if template['minio_path'] and self.minio_client:
                try:
                    template_content = self.minio_client.download_file(template['minio_path'])
                    template_text = self._extract_text_from_file(template_content, template['file_type'])
                except Exception as e:
                    log_.warning(f"提取模板文本失败: {str(e)}")
            
            if not self.ollama_model:
                raise APIException(ErrorCode.SYSTEM_ERROR, msg="AI模型未初始化")
            
            generated_text = self._generate_content_with_ai(template_text, user_input, template['name'])
            
            doc_data, file_ext = self._create_document_file(generated_text, output_format)
            
            minio_path = f"generated/{user_id}/{datetime.now().strftime('%Y%m%d%H%M%S')}_{template['name']}.{file_ext}"
            
            if self.minio_client:
                self.minio_client.upload_bytes(
                    data=doc_data,
                    object_name=minio_path,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_ext == 'docx' else 'application/pdf'
                )
            
            generated_doc = GeneratedDocument(
                template_id=template_id,
                template_name=template['name'],
                user_input=user_input,
                generated_content=generated_text,
                output_format=output_format,
                minio_path=minio_path,
                created_by=user_id
            )
            
            result = GeneratedDocumentDAO.create(generated_doc)
            return result
            
        except APIException:
            raise
        except Exception as e:
            log_.error(f"生成文档失败: {str(e)}", exc_info=True)
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"生成文档失败: {str(e)}")
    
    def get_generation_history(self, user_id, limit=20):
        """获取生成历史"""
        try:
            user_id = int(user_id)
            return GeneratedDocumentDAO.find_by_user_id(user_id, limit)
        except Exception as e:
            log_.error(f"获取生成历史失败: {str(e)}")
            raise APIException(ErrorCode.DATABASE_ERROR, msg=f"获取生成历史失败: {str(e)}")
    
    def download_generated_document(self, doc_id, user_id):
        """下载生成的文档"""
        try:
            user_id = int(user_id)
            doc_id = int(doc_id)
            
            doc = GeneratedDocumentDAO.find_by_id(doc_id)
            if not doc:
                raise APIException(ErrorCode.RESOURCE_NOT_FOUND, msg="文档不存在")
            
            if doc['created_by'] != user_id:
                raise APIException(ErrorCode.UNAUTHORIZED, msg="无权下载此文档")
            
            if not self.minio_client:
                raise APIException(ErrorCode.SYSTEM_ERROR, msg="存储服务未初始化")
            
            file_data = self.minio_client.download_file(doc['minio_path'])
            file_ext = 'docx' if doc['output_format'] == 'word' else 'pdf'
            filename = f"{doc['template_name']}.{file_ext}"
            
            return file_data, filename, file_ext
            
        except APIException:
            raise
        except Exception as e:
            log_.error(f"下载文档失败: {str(e)}", exc_info=True)
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"下载文档失败: {str(e)}")
    
    def _get_file_type(self, filename):
        """获取文件类型"""
        ext = os.path.splitext(filename)[1].lower()
        type_map = {
            '.doc': 'word',
            '.docx': 'word',
            '.pdf': 'pdf',
            '.txt': 'txt'
        }
        return type_map.get(ext, 'unknown')
    
    def _get_file_size(self, file_object):
        """获取文件大小"""
        try:
            pos = file_object.tell()
            file_object.seek(0, 2)
            size = file_object.tell()
            file_object.seek(pos)
            return size
        except:
            return 0
    
    def _get_content_type(self, file_type):
        """获取内容类型"""
        content_types = {
            'word': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'pdf': 'application/pdf',
            'txt': 'text/plain'
        }
        return content_types.get(file_type, 'application/octet-stream')
    
    def _extract_text_from_file(self, file_data, file_type):
        """从文件中提取文本"""
        try:
            if file_type == 'txt':
                return file_data.decode('utf-8') if isinstance(file_data, bytes) else str(file_data)
            elif file_type == 'word':
                return self._extract_from_word(file_data)
            elif file_type == 'pdf':
                return self._extract_from_pdf(file_data)
            else:
                return file_data.decode('utf-8') if isinstance(file_data, bytes) else str(file_data)
        except Exception as e:
            log_.error(f"提取文件文本失败: {str(e)}")
            return ""
    
    def _extract_from_word(self, file_data):
        """从Word文档提取文本"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_data))
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            return '\n'.join(text_parts)
        except Exception as e:
            log_.error(f"提取Word文本失败: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, file_data):
        """从PDF提取文本"""
        try:
            import fitz
            doc = fitz.open(stream=file_data, filetype="pdf")
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            return '\n'.join(text_parts)
        except Exception as e:
            log_.error(f"提取PDF文本失败: {str(e)}")
            return ""
    
    def _generate_content_with_ai(self, template_text, user_input, template_name):
        """使用AI生成文档内容"""
        try:
            prompt = f"""你是一个专业的文档生成助手。请根据以下信息生成一份完整的文档。

模板类型：{template_name}

模板格式参考：
{template_text[:2000]}

用户提供的个人信息：
{user_input}

请根据模板格式和用户提供的信息，生成一份完整、专业的文档内容。
要求：
1. 保持专业、正式的语言风格
2. 合理组织内容结构
3. 确保信息完整、准确
4. 直接输出文档内容，不要添加额外的解释说明

请直接输出生成的文档内容："""

            response = self.ollama_model.invoke(prompt)
            return response.strip()
            
        except Exception as e:
            log_.error(f"AI生成内容失败: {str(e)}")
            raise APIException(ErrorCode.MODEL_CALL_ERROR, msg=f"AI生成内容失败: {str(e)}")
    
    def _create_document_file(self, content, output_format):
        """创建文档文件"""
        try:
            if output_format == 'pdf':
                return self._create_pdf(content), 'pdf'
            else:
                return self._create_word(content), 'docx'
        except Exception as e:
            log_.error(f"创建文档文件失败: {str(e)}")
            raise APIException(ErrorCode.SYSTEM_ERROR, msg=f"创建文档文件失败: {str(e)}")
    
    def _create_word(self, content):
        """创建Word文档"""
        from docx import Document
        doc = Document()
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.read()
    
    def _create_pdf(self, content):
        """创建PDF文档"""
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import cm
        
        file_stream = io.BytesIO()
        doc = SimpleDocTemplate(file_stream, pagesize=A4,
                               rightMargin=2*cm, leftMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
        
        styles = getSampleStyleSheet()
        
        try:
            font_path = self._find_chinese_font()
            if font_path:
                pdfmetrics.registerFont(TTFont('Chinese', font_path))
                chinese_style = ParagraphStyle(
                    'Chinese',
                    parent=styles['Normal'],
                    fontName='Chinese',
                    fontSize=12,
                    leading=20
                )
                default_style = chinese_style
            else:
                default_style = styles['Normal']
        except:
            default_style = styles['Normal']
        
        story = []
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                story.append(Paragraph(para_text, default_style))
                story.append(Spacer(1, 0.3*cm))
        
        doc.build(story)
        file_stream.seek(0)
        return file_stream.read()
    
    def _find_chinese_font(self):
        """查找中文字体"""
        font_paths = [
            'C:/Windows/Fonts/simhei.ttf',
            'C:/Windows/Fonts/msyh.ttc',
            'C:/Windows/Fonts/simsun.ttc',
            '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        ]
        for path in font_paths:
            if os.path.exists(path):
                return path
        return None
